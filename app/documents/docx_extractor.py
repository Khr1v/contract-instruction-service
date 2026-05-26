from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.documents.base import DocumentExtractor
from app.llm.schemas import CanonicalDocument, PageOrSection, SourceFormat
from app.utils.text_utils import compact_whitespace


class DOCXExtractor(DocumentExtractor):
    def extract(self, file_path: str | Path, source_file_id: str, filename: str) -> CanonicalDocument:
        path = Path(file_path)
        document = Document(path)
        warnings = self._collect_warnings(document)

        blocks: list[PageOrSection] = []
        tables: list[str] = []
        text_parts: list[str] = []
        block_index = 1

        for block in self._iter_block_items(document):
            ref = f"section_{block_index}"
            if isinstance(block, Paragraph):
                text = compact_whitespace(block.text)
                if not text:
                    continue
                title = text[:80]
                blocks.append(PageOrSection(ref=ref, title=title, text=text, quality=1.0))
                text_parts.append(f"[{ref}]\n{text}")
                block_index += 1
            elif isinstance(block, Table):
                markdown = self._table_to_markdown(block)
                if not markdown.strip():
                    continue
                tables.append(markdown)
                blocks.append(
                    PageOrSection(
                        ref=ref,
                        title="Таблица",
                        text=markdown,
                        tables=[markdown],
                        quality=1.0,
                    )
                )
                text_parts.append(f"[{ref}]\n{markdown}")
                block_index += 1

        if not text_parts:
            warnings.append("DOCX не содержит извлекаемого текста.")

        return CanonicalDocument(
            source_file_id=source_file_id,
            filename=filename,
            source_format=SourceFormat.DOCX,
            document_text="\n\n".join(text_parts),
            pages_or_sections=blocks,
            tables=tables,
            metadata={
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
            },
            extraction_warnings=warnings,
            quality_score=1.0 if text_parts else 0.2,
        )

    def _iter_block_items(self, document: DocxDocument) -> Iterable[Paragraph | Table]:
        parent_elm = document.element.body
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield Table(child, document)

    def _table_to_markdown(self, table: Table) -> str:
        rows = [[compact_whitespace(cell.text.replace("\n", " ")) for cell in row.cells] for row in table.rows]
        if not rows:
            return ""
        max_columns = max(len(row) for row in rows)
        normalized = [row + [""] * (max_columns - len(row)) for row in rows]
        header = normalized[0]
        separator = ["---"] * max_columns
        body = normalized[1:]
        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(separator) + " |",
        ]
        for row in body:
            lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)

    def _collect_warnings(self, document: DocxDocument) -> list[str]:
        warnings: list[str] = []
        if document.inline_shapes:
            warnings.append("DOCX содержит изображения; текст на изображениях не извлекался.")

        xml = document.element.xml
        if "w:ins" in xml or "w:del" in xml:
            warnings.append("DOCX содержит признаки tracked changes; проверьте актуальность текста.")
        if "w:commentRangeStart" in xml or "comments.xml" in xml:
            warnings.append("DOCX содержит признаки комментариев; комментарии не интерпретировались как условия договора.")
        return warnings

