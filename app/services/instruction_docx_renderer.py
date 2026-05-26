from __future__ import annotations

import importlib.util
import re
import sys
from pathlib import Path
from types import ModuleType

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.table import _Cell

from app.config import PROJECT_ROOT
from contract_rules_agent.models import (
    ContactRow,
    CustomerInstructionData,
    GuaranteeLane,
    PlatformRow,
    StatusInforming,
)


class InstructionDocxRenderer:
    """Render generated Markdown instruction into a readable DOCX artifact."""

    def __init__(self, provided_renderer_path: str | Path | None = None) -> None:
        self.provided_renderer_path = Path(provided_renderer_path) if provided_renderer_path else None
        self._provided_renderer: ModuleType | None = None

    def render_template_instruction(
        self,
        *,
        renderer_data: dict[str, object],
        template_path: str | Path,
        output_path: str | Path,
    ) -> Path:
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        renderer = self._load_provided_renderer()
        data = self._build_customer_instruction_data(renderer_data)
        renderer.render_instruction_docx(Path(template_path), output, data)
        return output

    def render_template_markdown(
        self,
        *,
        renderer_data: dict[str, object],
        template_path: str | Path,
    ) -> str:
        renderer = self._load_provided_renderer()
        data = self._build_customer_instruction_data(renderer_data)
        return renderer.render_instruction_markdown(data, Path(template_path))

    def _load_provided_renderer(self) -> ModuleType:
        if self._provided_renderer is not None:
            return self._provided_renderer

        candidates = []
        if self.provided_renderer_path:
            candidates.append(self.provided_renderer_path)
        candidates.append(PROJECT_ROOT / "app" / "services" / "provided_renderer.py")

        for candidate in candidates:
            if not candidate.exists():
                continue
            spec = importlib.util.spec_from_file_location("provided_instruction_renderer", candidate)
            if spec is None or spec.loader is None:
                continue
            module = importlib.util.module_from_spec(spec)
            sys.modules[spec.name] = module
            spec.loader.exec_module(module)
            if not hasattr(module, "render_instruction_docx"):
                raise RuntimeError(f"Renderer has no render_instruction_docx: {candidate}")
            self._patch_provided_renderer(module)
            self._provided_renderer = module
            return module
        raise FileNotFoundError(
            "Provided renderer was not found. Set INSTRUCTION_RENDERER_PATH in .env "
            "or place renderer at app/services/provided_renderer.py."
        )

    def _patch_provided_renderer(self, module: ModuleType) -> None:
        original = getattr(module, "_replace_prefixed_line", None)
        if original is None or getattr(original, "_marshal_patched", False):
            return

        def patched_replace_prefixed_line(document: Document, prefix: str, value: str) -> None:
            original(document, prefix, value)
            if prefix != "1. Форма работы:":
                return
            paragraph = None
            for candidate in document.paragraphs:
                if candidate.text.strip().startswith(prefix):
                    paragraph = candidate
                    break
            if paragraph is None:
                return
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            if paragraph.runs:
                paragraph.runs[0].bold = True
                paragraph.runs[0].italic = True

        patched_replace_prefixed_line._marshal_patched = True
        module._replace_prefixed_line = patched_replace_prefixed_line

    def _build_customer_instruction_data(self, payload: dict[str, object]) -> CustomerInstructionData:
        return CustomerInstructionData(
            client_name=self._str(payload.get("client_name")),
            contract_legal_entity=self._str(payload.get("contract_legal_entity")),
            generated_date=self._str(payload.get("generated_date")),
            work_format=self._str(payload.get("work_format")),
            guarantee_lanes=[
                GuaranteeLane(
                    direction=self._str(item.get("direction")),
                    cost=self._str(item.get("cost")),
                    vehicles_count=self._str(item.get("vehicles_count")),
                    special_conditions=self._str(item.get("special_conditions")),
                )
                for item in self._dict_list(payload.get("guarantee_lanes"))
            ],
            guaranteed_application_rules=self._str_list(payload.get("guaranteed_application_rules")),
            spot_application_rules=self._str_list(payload.get("spot_application_rules")),
            platform_rows=[
                PlatformRow(
                    platform_name=self._str(item.get("platform_name")),
                    credentials=self._str(item.get("credentials")),
                    bidding_rules=self._str(item.get("bidding_rules")),
                    instruction_link=self._str(item.get("instruction_link")),
                )
                for item in self._dict_list(payload.get("platform_rows"))
            ],
            communication_rows=[
                ContactRow(
                    full_name=self._str(item.get("full_name")),
                    role=self._str(item.get("role")),
                    contacts=self._str(item.get("contacts")),
                    responsibility=self._str(item.get("responsibility")),
                )
                for item in self._dict_list(payload.get("communication_rows"))
            ],
            loading_requirements=self._str_list(payload.get("loading_requirements")),
            unloading_requirements=self._str_list(payload.get("unloading_requirements")),
            special_conditions=self._str_list(payload.get("special_conditions")),
            driver_briefing=self._str_list(payload.get("driver_briefing")),
            penalties=self._str_list(payload.get("penalties")),
            incident_actions=self._str_list(payload.get("incident_actions")),
            status_informing=StatusInforming(
                is_required=self._str(self._dict(payload.get("status_informing")).get("is_required")),
                frequency=self._str(self._dict(payload.get("status_informing")).get("frequency")),
                channels=self._str_list(self._dict(payload.get("status_informing")).get("channels")),
            ),
            payment_document_package=self._str_list(payload.get("payment_document_package")),
            payment_document_package_auto=self._str_list(payload.get("payment_document_package_auto")),
            payment_document_package_rail=self._str_list(payload.get("payment_document_package_rail")),
            document_format_requirements=self._str_list(payload.get("document_format_requirements")),
            copies_followed_by_originals=self._str(payload.get("copies_followed_by_originals")),
            edo_workflow=self._str(payload.get("edo_workflow")),
            client_document_contact=self._str(payload.get("client_document_contact")),
            executor_document_contact=self._str(payload.get("executor_document_contact")),
            copies_email=self._str(payload.get("copies_email")),
            originals_postal_address=self._str(payload.get("originals_postal_address")),
            payment_term=self._str(payload.get("payment_term")),
            payment_hold_condition=self._str(payload.get("payment_hold_condition")),
            tax_change_notification=self._str(payload.get("tax_change_notification")),
            client_payment_delay_penalty=self._str(payload.get("client_payment_delay_penalty")),
            open_questions=self._str_list(payload.get("open_questions")),
            extraction_notes=self._str_list(payload.get("extraction_notes")),
        )

    def _dict(self, value: object) -> dict[str, object]:
        return value if isinstance(value, dict) else {}

    def _dict_list(self, value: object) -> list[dict[str, object]]:
        if not isinstance(value, list):
            return []
        return [item for item in value if isinstance(item, dict)]

    def _str_list(self, value: object) -> list[str]:
        if value is None:
            return []
        if isinstance(value, list):
            return [self._str(item) for item in value if self._str(item)]
        text = self._str(value)
        return [text] if text else []

    def _str(self, value: object) -> str:
        if value is None:
            return ""
        return str(value).strip()

    def render_markdown_to_docx(
        self,
        markdown: str,
        output_path: str | Path,
        *,
        title: str = "ИНСТРУКЦИЯ ПО РАБОТЕ С КЛИЕНТОМ",
    ) -> Path:
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        self._configure_document(document)
        self._configure_styles(document)
        self._write_markdown(document, markdown, title=title)
        document.save(output)
        return output

    def _configure_document(self, document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)
        section.start_type = WD_SECTION_START.NEW_PAGE

    def _configure_styles(self, document: Document) -> None:
        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.08

        for name, size, color in [
            ("Heading 1", 16, RGBColor(31, 78, 121)),
            ("Heading 2", 13, RGBColor(31, 78, 121)),
            ("Heading 3", 11, RGBColor(31, 78, 121)),
        ]:
            style = document.styles[name]
            style.font.name = "Arial"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = color
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(6)

    def _write_markdown(self, document: Document, markdown: str, *, title: str) -> None:
        lines = markdown.splitlines()
        index = 0
        wrote_title = False

        while index < len(lines):
            line = lines[index].rstrip()
            stripped = line.strip()
            if not stripped:
                index += 1
                continue
            if stripped == "---":
                index += 1
                continue
            if self._is_table_start(lines, index):
                table_lines: list[str] = []
                while index < len(lines) and lines[index].strip().startswith("|"):
                    table_lines.append(lines[index].strip())
                    index += 1
                self._add_markdown_table(document, table_lines)
                continue

            heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
            if heading:
                level = min(len(heading.group(1)), 3)
                text = self._clean_inline_markdown(heading.group(2))
                if not wrote_title and "ИНСТРУКЦИЯ ПО РАБОТЕ" in text.upper():
                    self._add_title(document, title)
                    wrote_title = True
                else:
                    document.add_paragraph(text, style=f"Heading {level}")
                index += 1
                continue

            if not wrote_title and "ИНСТРУКЦИЯ ПО РАБОТЕ" in stripped.upper():
                self._add_title(document, title)
                wrote_title = True
                index += 1
                continue

            if stripped.startswith(("- ", "* ", "• ")):
                paragraph = document.add_paragraph(style="List Bullet")
                self._add_inline_runs(paragraph, stripped[2:].strip())
                index += 1
                continue

            numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
            if numbered and len(stripped) < 220:
                paragraph = document.add_paragraph(style="List Number")
                self._add_inline_runs(paragraph, numbered.group(1).strip())
                index += 1
                continue

            paragraph = document.add_paragraph()
            self._add_inline_runs(paragraph, stripped)
            index += 1

    def _add_title(self, document: Document, title: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(title)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(31, 78, 121)

    def _is_table_start(self, lines: list[str], index: int) -> bool:
        return (
            index + 1 < len(lines)
            and lines[index].strip().startswith("|")
            and re.match(r"^\|\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", lines[index + 1].strip()) is not None
        )

    def _add_markdown_table(self, document: Document, table_lines: list[str]) -> None:
        rows = [self._split_markdown_row(line) for line in table_lines]
        if len(rows) < 2:
            return
        rows = [rows[0], *rows[2:]]
        column_count = max(len(row) for row in rows)
        rows = [row + [""] * (column_count - len(row)) for row in rows]

        table = document.add_table(rows=len(rows), cols=column_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        table.autofit = True

        for row_index, row in enumerate(rows):
            for col_index, value in enumerate(row):
                cell = table.cell(row_index, col_index)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                self._set_cell_text(cell, value, bold=row_index == 0)
        document.add_paragraph()

    def _split_markdown_row(self, line: str) -> list[str]:
        text = line.strip().strip("|")
        cells: list[str] = []
        current: list[str] = []
        escaped = False
        for char in text:
            if escaped:
                current.append(char)
                escaped = False
            elif char == "\\":
                escaped = True
            elif char == "|":
                cells.append("".join(current).strip())
                current = []
            else:
                current.append(char)
        cells.append("".join(current).strip())
        return cells

    def _set_cell_text(self, cell: _Cell, value: str, *, bold: bool) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        parts = re.split(r"<br\s*/?>", value)
        for part_index, part in enumerate(parts):
            if part_index:
                paragraph.add_run().add_break()
            self._add_inline_runs(paragraph, part, force_bold=bold)

    def _add_inline_runs(self, paragraph, text: str, *, force_bold: bool = False) -> None:
        chunks = re.split(r"(\*\*.*?\*\*)", text)
        for chunk in chunks:
            if not chunk:
                continue
            bold = force_bold
            if chunk.startswith("**") and chunk.endswith("**"):
                chunk = chunk[2:-2]
                bold = True
            run = paragraph.add_run(self._clean_inline_markdown(chunk))
            run.bold = bold
            run.font.name = "Arial"
            run.font.size = Pt(10)

    def _clean_inline_markdown(self, text: str) -> str:
        text = text.replace("<br>", "\n")
        text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        return text.strip()
