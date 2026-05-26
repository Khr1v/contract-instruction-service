from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

from contract_rules_agent.models import ContactRow, CustomerInstructionData, PlatformRow


SECTION_HEADING_RE = re.compile(r"^\d+\.\s")
MISSING_MARKER = "Требует уточнения по договору."
BODY_STYLE = "Body Text"
LIST_STYLE = "List Paragraph"
TABLE_STYLE = "Table Paragraph"
ACCENT_BLUE = RGBColor(141, 179, 255)
PENALTY_HIGHLIGHT_PATTERNS = (
    re.compile(r"штраф(?:а)?\s+в\s+размере\s+\d+\s?%", flags=re.IGNORECASE),
    re.compile(r"\b\d+\s?%\s+от\s+ставки\b", flags=re.IGNORECASE),
    re.compile(r"\b\d+\s?%\s+от\s+стоимости\s+перевозки\b", flags=re.IGNORECASE),
    re.compile(
        r"\b\d[\d\s]*(?:,\d+)?\s*руб(?:л(?:ей|я)|\.)\s+за\s+каждый\s+документ\b",
        flags=re.IGNORECASE,
    ),
    re.compile(
        r"штраф(?:а)?\s+в\s+размере\s+\d[\d\s]*(?:,\d+)?\s*руб(?:л(?:ей|я)|\.)",
        flags=re.IGNORECASE,
    ),
    re.compile(r"Срывом считается[^.]+(?:\.)?", flags=re.IGNORECASE),
)
LOADING_TRACTOR_LINE = "Под погрузку должно подаваться ТС в составе тягача и полуприцепа"
LOADING_PNEUMO_LINE = "Тягач и полуприцеп должны быть оборудованы исправной пневматической подвеской"
LOADING_TRACTOR_PNEUMO_LINE = (
    "Под погрузку должно подаваться ТС в составе тягача и полуприцепа с исправной пневматической подвеской"
)
LOADING_TECH_LINE = (
    "ТС подается под погрузку в технически исправном, чистом и сухом состоянии: "
    "герметичный кузов, целый тент, стандартная обрешетка, без посторонних предметов"
)
LOADING_FIXTURES_LINE = (
    "В ТС должны быть исправные механизмы крепления груза; пол и стены кузова должны быть ровными, "
    "чтобы исключить повреждение груза"
)
LOADING_TECH_FIXTURES_LINE = (
    "ТС подается под погрузку в технически исправном, чистом и сухом состоянии; "
    "в ТС должны быть исправные механизмы крепления груза"
)
DOCUMENT_TRUST_REQUIREMENT_LINE = (
    "Экспедитор обязан оформить и передать Клиенту доверенность на получение груза "
    "и перевозку водителем транспортного средства; оригинал доверенности направляется дополнительно"
)


@dataclass(slots=True)
class ParagraphBlock:
    text: str
    style_name: str = BODY_STYLE
    bold: bool = False
    italic: bool = False
    bullet: bool = False
    accent: bool = False
    penalty_highlight: bool = False


@dataclass(slots=True)
class PenaltyGroup:
    key: str
    title: str
    intro: str = ""
    items: list[str] = field(default_factory=list)


def render_instruction_markdown(data: CustomerInstructionData, template_path: Path) -> str:
    loading_items = _prepare_loading_requirements(data.loading_requirements)
    special_items = _prepare_special_conditions(data.special_conditions)
    incident_items = _prepare_incident_actions(data.incident_actions)
    document_requirement_items = _prepare_document_requirements(data.document_format_requirements)
    status_frequency = _render_status_frequency(data)
    guaranteed_blocks = _build_application_blocks(data.guaranteed_application_rules)
    spot_blocks = _build_application_blocks(data.spot_application_rules)
    loading_blocks = _build_requirement_blocks(
        loading_items,
        intro_text="Обязательные требования к подвижному составу:",
    )
    unloading_blocks = _build_requirement_blocks(data.unloading_requirements)
    special_blocks = _build_requirement_blocks(special_items)
    driver_blocks = _build_requirement_blocks(data.driver_briefing)
    penalty_blocks = _build_penalty_blocks(data.penalties)
    incident_blocks = _build_incident_blocks(incident_items)

    return (
        "# Инструкция по работе с клиентом\n\n"
        f"Клиент: {data.client_name or MISSING_MARKER}\n"
        f"Юридическое лицо по договору: {data.contract_legal_entity or MISSING_MARKER}\n"
        f"Дата создания/обновления документа: {data.generated_date or MISSING_MARKER}\n"
        f"Шаблон: {template_path.name}\n\n"
        "## 1. Форма работы\n\n"
        f"Режим: {data.work_format or MISSING_MARKER}\n\n"
        "Гарантированные заявки:\n"
        f"{_format_markdown_blocks(guaranteed_blocks)}\n\n"
        "Спотовые заявки:\n"
        f"{_format_markdown_blocks(spot_blocks)}\n\n"
        "## 2. Матрица коммуникаций клиента\n\n"
        f"{_format_contacts(data)}\n\n"
        "## 5. Требования на погрузке\n\n"
        f"{_format_markdown_blocks(loading_blocks)}\n\n"
        "## 6. Требования на выгрузке\n\n"
        f"{_format_markdown_blocks(unloading_blocks)}\n\n"
        "## 7. Особые условия\n\n"
        f"{_format_markdown_blocks(special_blocks)}\n\n"
        "## 8. Инструктаж для водителя\n\n"
        f"{_format_markdown_blocks(driver_blocks)}\n\n"
        "## 9. Штрафы\n\n"
        f"{_format_markdown_blocks(penalty_blocks)}\n\n"
        "## 10. Действия логиста при проблемных ситуациях\n\n"
        f"{_format_markdown_blocks(incident_blocks)}\n\n"
        "Информирование клиента о статусе перевозки\n\n"
        f"- Требуется: {data.status_informing.is_required or MISSING_MARKER}\n"
        f"- Частота: {status_frequency or MISSING_MARKER}\n"
        f"- Каналы: {', '.join(data.status_informing.channels) if data.status_informing.channels else MISSING_MARKER}\n\n"
        "## 11. Документооборот\n\n"
        f"- Комплект документов для оплаты:\n{_format_indented_lines(_build_payment_package_lines(data))}\n"
        f"- Требования к оформлению:\n{_format_indented_lines(_build_bulleted_lines(document_requirement_items))}\n"
        f"- Работа через ЭДО: {_condense_edo_workflow(data.edo_workflow) or MISSING_MARKER}\n"
        f"- Контакт клиента по документообороту: {data.client_document_contact or MISSING_MARKER}\n"
        f"- Контакт исполнителя по документообороту: {data.executor_document_contact or MISSING_MARKER}\n"
        f"- Email для копий: {data.copies_email or MISSING_MARKER}\n"
        f"- Почтовый адрес для оригиналов: {data.originals_postal_address or MISSING_MARKER}\n\n"
        "## 12. Условия оплаты\n\n"
        f"{_format_markdown_blocks(_build_payment_blocks(data))}\n\n"
        "## Что проверить вручную\n\n"
        f"{_format_optional_list(data.open_questions, 'Нет обязательных ручных замечаний по извлечению.')}\n\n"
        "## Технические заметки\n\n"
        f"{_format_optional_list(data.extraction_notes, 'Нет технических замечаний.')}\n"
    )


def render_instruction_docx(
    template_path: Path,
    output_path: Path,
    data: CustomerInstructionData,
) -> None:
    document = Document(str(template_path))
    _ensure_instruction_styles(document)
    loading_items = _prepare_loading_requirements(data.loading_requirements)
    special_items = _prepare_special_conditions(data.special_conditions)
    incident_items = _prepare_incident_actions(data.incident_actions)

    _replace_prefixed_line(document, "Клиент:", data.client_name or MISSING_MARKER)
    _replace_prefixed_line(
        document,
        "Юридическое лицо по договору:",
        data.contract_legal_entity or MISSING_MARKER,
    )
    _replace_prefixed_line(
        document,
        "Дата создания/обновления документа:",
        data.generated_date or MISSING_MARKER,
    )
    _replace_prefixed_line(
        document,
        "1. Форма работы:",
        data.work_format or "Гарантии / Спот",
    )

    _replace_section_body(
        document,
        "Гарантированные заявки",
        _build_application_blocks(data.guaranteed_application_rules),
        stop_text="Спотовые заявки",
    )
    _replace_section_body(
        document,
        "Спотовые заявки",
        _build_application_blocks(data.spot_application_rules),
        stop_text="2. Матрица коммуникаций Клиента",
    )
    _replace_section_body(
        document,
        "5. Требования на погрузке",
        _build_requirement_blocks(
            loading_items,
            intro_text="Обязательные требования к подвижному составу:",
        ),
        stop_text="6. Требования на выгрузке",
    )
    _replace_section_body(
        document,
        "6. Требования на выгрузке",
        _build_requirement_blocks(data.unloading_requirements),
        stop_text="7. Особые условия",
    )
    _replace_section_body(
        document,
        "7. Особые условия",
        _build_requirement_blocks(special_items),
        stop_text="8. Инструктаж для водителя",
    )
    _replace_section_body(
        document,
        "8. Инструктаж для водителя",
        _build_requirement_blocks(data.driver_briefing),
        stop_text="9. Штрафы",
    )
    _replace_section_body(
        document,
        "9. Штрафы",
        _build_penalty_blocks(data.penalties),
        stop_text="10. Действия логиста при проблемных ситуациях",
    )
    _replace_section_body(
        document,
        "10. Действия логиста при проблемных ситуациях",
        _build_incident_blocks(incident_items),
        stop_text="11. Информирование Клиента о статусе перевозки",
    )
    _replace_heading_text(
        document,
        "11. Информирование Клиента о статусе перевозки",
        "Информирование Клиента о статусе перевозки",
    )

    if document.tables:
        _fill_guarantee_table(document.tables[0], data)
    if len(document.tables) > 1:
        _fill_platform_table(document.tables[1], data)
    if len(document.tables) > 2:
        _fill_communication_table(document.tables[2], data)
    if len(document.tables) > 3:
        _fill_status_table(document.tables[3], data)
    if len(document.tables) > 4:
        _fill_document_flow_table(document.tables[4], data)

    _append_payment_section(document, data)
    _apply_section_spacing(document)
    _apply_accent_styling(document)

    document.save(str(output_path))


def _ensure_instruction_styles(document: Document) -> None:
    normal_style = document.styles["Normal"]

    body_style = document.styles[BODY_STYLE] if BODY_STYLE in document.styles else document.styles.add_style(
        BODY_STYLE,
        WD_STYLE_TYPE.PARAGRAPH,
    )
    body_style.base_style = normal_style
    body_style.font.name = "Times New Roman"
    body_style.font.size = Pt(12)
    body_style.paragraph_format.line_spacing = 1.15
    body_style.paragraph_format.first_line_indent = Pt(21)
    body_style.paragraph_format.space_before = Pt(0)
    body_style.paragraph_format.space_after = Pt(0)

    list_style = document.styles[LIST_STYLE] if LIST_STYLE in document.styles else document.styles.add_style(
        LIST_STYLE,
        WD_STYLE_TYPE.PARAGRAPH,
    )
    list_style.base_style = body_style
    list_style.font.name = "Times New Roman"
    list_style.font.size = Pt(12)
    list_style.paragraph_format.line_spacing = 1.1
    list_style.paragraph_format.left_indent = Pt(21)
    list_style.paragraph_format.first_line_indent = Pt(0)
    list_style.paragraph_format.space_before = Pt(1.5)
    list_style.paragraph_format.space_after = Pt(0)

    table_style = document.styles[TABLE_STYLE] if TABLE_STYLE in document.styles else document.styles.add_style(
        TABLE_STYLE,
        WD_STYLE_TYPE.PARAGRAPH,
    )
    table_style.base_style = normal_style
    table_style.font.name = "Times New Roman"
    table_style.font.size = Pt(10.5)
    table_style.paragraph_format.line_spacing = 1.0
    table_style.paragraph_format.first_line_indent = Pt(0)
    table_style.paragraph_format.left_indent = Pt(0)
    table_style.paragraph_format.space_before = Pt(0)
    table_style.paragraph_format.space_after = Pt(0)


def _replace_prefixed_line(
    document: Document,
    prefix: str,
    value: str,
) -> None:
    paragraph = _find_paragraph(document, prefix)
    if paragraph is None:
        return

    _clear_paragraph_content(paragraph)
    prefix_run = paragraph.add_run(prefix)
    prefix_run.bold = True
    prefix_run.italic = True
    prefix_run.font.name = "Times New Roman"
    prefix_run.font.size = Pt(14)

    spacer = paragraph.add_run(" ")
    spacer.font.name = "Times New Roman"
    spacer.font.size = Pt(14)

    value_run = paragraph.add_run(value)
    value_run.font.name = "Times New Roman"
    value_run.font.size = Pt(14)
    value_run.font.color.rgb = ACCENT_BLUE


def _replace_section_body(
    document: Document,
    heading_text: str,
    content_blocks: list[ParagraphBlock],
    stop_text: str,
) -> None:
    paragraphs = document.paragraphs
    heading_index = _find_paragraph_index(paragraphs, heading_text)
    if heading_index is None:
        return

    body_indices = _collect_body_indices(paragraphs, heading_index, stop_text)
    normalized_blocks = content_blocks or [ParagraphBlock(MISSING_MARKER)]

    if not body_indices:
        anchor = paragraphs[heading_index]
        inserted = anchor
        for block in normalized_blocks:
            inserted = _insert_paragraph_after(inserted, block)
        return

    for offset, body_index in enumerate(body_indices):
        if offset < len(normalized_blocks):
            _populate_paragraph(paragraphs[body_index], normalized_blocks[offset])
        else:
            _clear_paragraph_content(paragraphs[body_index])

    if len(normalized_blocks) > len(body_indices):
        anchor = paragraphs[body_indices[-1]]
        for block in normalized_blocks[len(body_indices):]:
            anchor = _insert_paragraph_after(anchor, block)


def _fill_guarantee_table(table, data: CustomerInstructionData) -> None:
    rows = data.guarantee_lanes or [None]
    _ensure_table_rows(table, len(rows) + 1)
    for row_index in range(1, len(table.rows)):
        cells = table.rows[row_index].cells
        if row_index - 1 < len(rows) and rows[row_index - 1] is not None:
            lane = rows[row_index - 1]
            _set_cell_lines(cells[0], [lane.direction or MISSING_MARKER], style_name=TABLE_STYLE)
            _set_cell_lines(cells[1], [lane.cost or ""], style_name=TABLE_STYLE)
            _set_cell_lines(cells[2], [lane.vehicles_count or ""], style_name=TABLE_STYLE)
            _set_cell_lines(cells[3], [lane.special_conditions or ""], style_name=TABLE_STYLE)
        elif row_index == 1:
            _set_cell_lines(cells[0], [MISSING_MARKER], style_name=TABLE_STYLE)
            _set_cell_lines(cells[1], [""], style_name=TABLE_STYLE)
            _set_cell_lines(cells[2], [""], style_name=TABLE_STYLE)
            _set_cell_lines(cells[3], [""], style_name=TABLE_STYLE)
        else:
            _set_cell_lines(cells[0], [""], style_name=TABLE_STYLE)
            _set_cell_lines(cells[1], [""], style_name=TABLE_STYLE)
            _set_cell_lines(cells[2], [""], style_name=TABLE_STYLE)
            _set_cell_lines(cells[3], [""], style_name=TABLE_STYLE)


def _fill_platform_table(table, data: CustomerInstructionData) -> None:
    rows = data.platform_rows or [PlatformRow(platform_name=MISSING_MARKER)]
    _ensure_table_rows(table, len(rows) + 1)
    for row_index in range(1, len(table.rows)):
        cells = table.rows[row_index].cells
        if row_index - 1 < len(rows):
            row = rows[row_index - 1]
            platform_name = _normalize_platform_name(row)
            _set_cell_lines(cells[0], [platform_name], style_name=TABLE_STYLE)
            _set_cell_lines(cells[1], [row.credentials or ""], style_name=TABLE_STYLE)
            _set_cell_lines(cells[2], [row.bidding_rules or ""], style_name=TABLE_STYLE)
            _set_cell_lines(cells[3], [row.instruction_link or ""], style_name=TABLE_STYLE)
        elif row_index == 1 and not data.platform_rows:
            _set_cell_lines(cells[0], [MISSING_MARKER], style_name=TABLE_STYLE)


def _fill_communication_table(table, data: CustomerInstructionData) -> None:
    template_slots = list(range(1, len(table.rows)))
    assigned_rows: set[int] = set()

    for row in data.communication_rows:
        target_index = _select_communication_row(table, row, template_slots, assigned_rows)
        if target_index is None:
            if template_slots:
                target_index = template_slots[-1]
            else:
                table.add_row()
                target_index = len(table.rows) - 1

        assigned_rows.add(target_index)
        cells = table.rows[target_index].cells
        display_name = _contact_display_name(row)
        display_role = _contact_display_role(row)
        _set_cell_lines(
            cells[0],
            [display_name],
            style_name=TABLE_STYLE,
        )
        _set_cell_lines(cells[1], [display_role], style_name=TABLE_STYLE)
        _set_cell_lines(cells[2], [row.contacts or ""], style_name=TABLE_STYLE)
        _set_cell_lines(cells[3], [row.responsibility or ""], style_name=TABLE_STYLE)

    if not data.communication_rows and len(table.rows) > 1:
        cells = table.rows[1].cells
        _set_cell_lines(cells[0], [MISSING_MARKER], style_name=TABLE_STYLE)


def _fill_status_table(table, data: CustomerInstructionData) -> None:
    rows = table.rows
    _set_cell_lines(rows[0].cells[1], [data.status_informing.is_required or MISSING_MARKER], style_name=TABLE_STYLE)
    _set_cell_lines(rows[1].cells[1], [_render_status_frequency(data) or MISSING_MARKER], style_name=TABLE_STYLE)
    _set_cell_lines(
        rows[2].cells[1],
        data.status_informing.channels or [MISSING_MARKER],
        style_name=TABLE_STYLE,
    )


def _fill_document_flow_table(table, data: CustomerInstructionData) -> None:
    values = [
        _build_payment_package_blocks(data),
        _build_document_requirement_blocks(data.document_format_requirements),
        [_table_text_block("См. п. 12 «Условия оплаты».")],
        [data.copies_followed_by_originals or MISSING_MARKER],
        [_condense_edo_workflow(data.edo_workflow) or MISSING_MARKER],
        [data.client_document_contact or MISSING_MARKER],
        [data.executor_document_contact or MISSING_MARKER],
        [data.copies_email or MISSING_MARKER],
        [data.originals_postal_address or MISSING_MARKER],
    ]
    _ensure_table_rows(table, len(values))
    for row_index, value_lines in enumerate(values):
        if row_index >= len(table.rows):
            break
        _set_cell_lines(table.rows[row_index].cells[1], value_lines, style_name=TABLE_STYLE)


def _ensure_table_rows(table, required_rows: int) -> None:
    while len(table.rows) < required_rows:
        table.add_row()


def _find_paragraph(document: Document, startswith: str) -> Paragraph | None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            return paragraph
    return None


def _find_paragraph_index(paragraphs: list[Paragraph], text: str) -> int | None:
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == text:
            return index
    return None


def _collect_body_indices(
    paragraphs: list[Paragraph],
    heading_index: int,
    stop_text: str,
) -> list[int]:
    result: list[int] = []
    for index in range(heading_index + 1, len(paragraphs)):
        text = paragraphs[index].text.strip()
        if text == stop_text:
            break
        if text and SECTION_HEADING_RE.match(text):
            break
        result.append(index)
    return result


def _insert_paragraph_after(
    paragraph: Paragraph,
    block: ParagraphBlock,
) -> Paragraph:
    new_paragraph_element = OxmlElement("w:p")
    paragraph._p.addnext(new_paragraph_element)
    new_paragraph = Paragraph(new_paragraph_element, paragraph._parent)
    _populate_paragraph(new_paragraph, block)
    return new_paragraph


def _populate_paragraph(paragraph: Paragraph, block: ParagraphBlock) -> None:
    _clear_paragraph_content(paragraph)
    _set_paragraph_style(paragraph, block.style_name)
    if block.bullet:
        bullet_run = paragraph.add_run("• ")
        bullet_run.font.name = "Times New Roman"
        bullet_run.font.size = Pt(12 if block.style_name != TABLE_STYLE else 10.5)
        bullet_run.bold = True
    font_size = Pt(12 if block.style_name != TABLE_STYLE else 10.5)
    text_runs = _segment_penalty_text(block.text) if block.penalty_highlight else [(block.text, False)]
    for text_part, is_highlighted in text_runs:
        run = paragraph.add_run(text_part)
        run.font.name = "Times New Roman"
        run.font.size = font_size
        if block.bold:
            run.bold = True
        if block.italic:
            run.italic = True
        if block.accent:
            run.font.color.rgb = ACCENT_BLUE
        if is_highlighted:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _set_paragraph_style(paragraph: Paragraph, style_name: str) -> None:
    try:
        paragraph.style = style_name
    except KeyError:
        paragraph.style = "Normal"


def _clear_paragraph_content(paragraph: Paragraph) -> None:
    paragraph_element = paragraph._p
    for child in list(paragraph_element):
        if child.tag != qn("w:pPr"):
            paragraph_element.remove(child)


def _set_cell_lines(cell, lines: list[str], style_name: str = BODY_STYLE) -> None:
    normalized_lines = lines or [""]
    cell.text = ""
    first_paragraph = cell.paragraphs[0]
    first_block = _normalize_cell_block(normalized_lines[0], style_name)
    _populate_paragraph(first_paragraph, first_block)
    for line in normalized_lines[1:]:
        paragraph = cell.add_paragraph()
        _populate_paragraph(paragraph, _normalize_cell_block(line, style_name))


def _build_application_blocks(items: list[str]) -> list[ParagraphBlock]:
    if not items:
        return [ParagraphBlock(MISSING_MARKER)]
    if len(items) == 1:
        return [ParagraphBlock(_ensure_sentence(items[0]), style_name=BODY_STYLE)]
    return [
        ParagraphBlock(_ensure_sentence(item), style_name=LIST_STYLE, bullet=True)
        for item in items
    ]


def _build_requirement_blocks(
    items: list[str],
    intro_text: str | None = None,
) -> list[ParagraphBlock]:
    if not items:
        return [ParagraphBlock(MISSING_MARKER)]

    narrative_items: list[str] = []
    compact_items: list[str] = []
    for item in items:
        cleaned = _clean_block_text(item)
        if _looks_like_compact_requirement(cleaned):
            compact_items.append(cleaned)
        else:
            narrative_items.append(cleaned)

    blocks = [
        ParagraphBlock(_ensure_sentence(item), style_name=BODY_STYLE)
        for item in narrative_items
    ]

    if compact_items:
        if intro_text and not any(_is_intro_like(item, intro_text) for item in narrative_items):
            blocks.append(ParagraphBlock(intro_text, style_name=BODY_STYLE))
        for index, item in enumerate(compact_items):
            blocks.append(
                ParagraphBlock(
                    _ensure_compact_requirement(item, is_last=index == len(compact_items) - 1),
                    style_name=LIST_STYLE,
                    bullet=True,
                )
            )

    return blocks or [ParagraphBlock(MISSING_MARKER)]


def _build_penalty_blocks(items: list[str]) -> list[ParagraphBlock]:
    if not items:
        return [ParagraphBlock(MISSING_MARKER)]

    expanded_items: list[str] = []
    for item in items:
        expanded_items.extend(_expand_penalty_item(item))

    grouped_items = _group_penalty_items(expanded_items)
    blocks: list[ParagraphBlock] = []
    if grouped_items:
        for index, group in enumerate(grouped_items, start=1):
            rendered_items = _prepare_penalty_group_render_items(group)
            blocks.append(
                ParagraphBlock(
                    f"{index}. {group.title}",
                    style_name=BODY_STYLE,
                    bold=True,
                )
            )
            if group.intro:
                intro_text = _condense_penalty_item(group.intro, group.key)
                if not intro_text:
                    intro_text = group.intro
                if intro_text and intro_text != group.title:
                    blocks.append(
                        ParagraphBlock(
                            _ensure_sentence(intro_text),
                            style_name=BODY_STYLE,
                        )
                    )

            if len(rendered_items) == 1 and not group.intro:
                blocks.append(
                    ParagraphBlock(
                        _ensure_sentence(rendered_items[0]),
                        style_name=BODY_STYLE,
                        penalty_highlight=True,
                    )
                )
                continue

            for condensed_item in rendered_items:
                blocks.append(
                    ParagraphBlock(
                        _ensure_sentence(condensed_item),
                        style_name=LIST_STYLE,
                        bullet=True,
                        penalty_highlight=True,
                    )
                )
    else:
        for item in expanded_items:
            cleaned = _clean_block_text(item)
            if cleaned.endswith(":") or "в размере:" in cleaned.lower():
                blocks.append(
                    ParagraphBlock(
                        _ensure_sentence(cleaned),
                        style_name=BODY_STYLE,
                        bold=True,
                        penalty_highlight=True,
                    )
                )
                continue
            blocks.append(
                ParagraphBlock(
                    _ensure_sentence(cleaned),
                    style_name=LIST_STYLE,
                    bullet=True,
                    penalty_highlight=True,
                )
            )

    return blocks or [ParagraphBlock(MISSING_MARKER)]


def _group_penalty_items(items: list[str]) -> list[PenaltyGroup]:
    groups: dict[str, PenaltyGroup] = {}
    order: list[str] = []

    for item in items:
        clean = _clean_block_text(item)
        if not clean:
            continue

        group_key, group_title = _classify_penalty_group(clean)
        if group_key not in groups:
            groups[group_key] = PenaltyGroup(key=group_key, title=group_title)
            order.append(group_key)

        group = groups[group_key]
        if _is_penalty_intro(clean, group_key) and not group.intro:
            group.intro = clean
        else:
            group.items.append(clean)

    return [groups[key] for key in order if groups[key].intro or groups[key].items]


def _classify_penalty_group(text: str) -> tuple[str, str]:
    lowered = text.casefold()

    if "срыв рейса" in lowered or (
        "отказ" in lowered and "подтвержден" in lowered and "поручени" in lowered
    ):
        return ("trip_failure", "Срыв рейса")

    if (
        "опоздани" in lowered
        and ("погрузк" in lowered or "выгрузк" in lowered or "ставк" in lowered)
    ) or "просрочки в доставке груза" in lowered or any(
        marker in lowered
        for marker in (
            "срок подачи транспортного средства",
            "задержки подачи",
            "неподачу транспортного средства",
            "неподача транспортного средства",
            "сверхнорматив",
            "за каждый час простоя",
            "сроков вывоза транспортного средства",
            "сроков вывоза транспортного средства на выгрузку",
            "срыв подачи транспортного средства",
            "срыв подачи тс",
            "срыв подачи",
        )
    ) or (
        "от стоимости перевозки" in lowered
        and any(marker in lowered for marker in ("сутки опоздания", "вторых суток", "каждые сутки"))
    ):
        return ("delay", "Опоздание на погрузку/выгрузку и просрочка доставки")

    if (
        (
            any(marker in lowered for marker in ("gps", "системы контроля"))
            and any(marker in lowered for marker in ("водител", "информац", "статус"))
        )
        or (
            "водител" in lowered
            and "информац" in lowered
            and any(marker in lowered for marker in ("не предостав", "нарушает срок", "недостовер"))
        )
    ):
        return ("gps", "Нарушения по GPS")

    document_penalty_markers = (
        "отчетных документ",
        "полного комплекта",
        "предоставления отчетных документов",
        "некорректно оформлен",
        "корректно оформлен",
        "акт оказан",
        "счет-фактур",
        "универсальн передаточн",
    )
    if (
        (
            any(marker in lowered for marker in document_penalty_markers)
            or re.search(r"\bупд\b", lowered)
        )
        and any(marker in lowered for marker in ("штраф", "неустойк", "отказаться от договора"))
    ):
        return ("documents", "Документы")

    return ("other", "Прочие штрафные санкции")


def _is_penalty_intro(text: str, group_key: str) -> bool:
    lowered = text.casefold()
    if group_key == "delay":
        return text.endswith(":") or "в размере:" in lowered
    return False


def _prepare_penalty_group_render_items(group: PenaltyGroup) -> list[str]:
    items: list[str] = []
    seen: set[str] = set()
    for item in group.items:
        condensed = _clean_block_text(_condense_penalty_item(item, group.key) or item)
        if not condensed:
            continue
        key = condensed.casefold()
        if key in seen:
            continue
        seen.add(key)
        items.append(condensed)

    if group.key == "delay":
        network_item = next(
            (item for item in items if "для сетевых клиентов" in item.casefold()),
            "",
        )
        other_item = next(
            (
                item
                for item in items
                if "по прочим грузополучателям" in item.casefold()
                or "для остальных" in item.casefold()
            ),
            "",
        )
        if network_item and other_item:
            merged_items: list[str] = []
            merged_added = False
            for item in items:
                lowered = item.casefold()
                if lowered == network_item.casefold() or lowered == other_item.casefold():
                    if not merged_added:
                        merged_items.append("За срыв подачи ТС: 50% для сетевых клиентов, 20% — для остальных")
                        merged_added = True
                    continue
                if lowered == "50% от стоимости услуг за неподачу транспортного средства в срок":
                    continue
                merged_items.append(item)
            items = merged_items

    if group.key == "delay":
        items.sort(key=_delay_penalty_sort_key)
    return items


def _delay_penalty_sort_key(text: str) -> tuple[int, str]:
    lowered = text.casefold()
    if "10%" in lowered or "10 %" in lowered:
        return (0, lowered)
    if "25%" in lowered or "25 %" in lowered:
        return (1, lowered)
    if "50%" in lowered or "50 %" in lowered:
        return (2, lowered)
    if "20%" in lowered or "20 %" in lowered:
        return (3, lowered)
    return (9, lowered)


def _prepare_loading_requirements(items: list[str]) -> list[str]:
    prepared = _prepare_section_items(items, _condense_loading_requirement)
    return _merge_loading_requirements(prepared)


def _prepare_document_requirements(items: list[str]) -> list[str]:
    return _prepare_section_items(items, _condense_document_requirement)


def _prepare_special_conditions(items: list[str]) -> list[str]:
    return _prepare_section_items(items, _condense_special_condition)


def _prepare_incident_actions(items: list[str]) -> list[str]:
    return _prepare_section_items(items, _condense_incident_action)


def _prepare_section_items(
    items: list[str],
    transformer: Callable[[str], str],
) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for item in items:
        transformed = _clean_block_text(transformer(item))
        if not transformed:
            continue
        key = transformed.casefold()
        if key in seen:
            continue
        seen.add(key)
        result.append(transformed)
    return result


def _condense_loading_requirement(text: str) -> str:
    clean = _clean_block_text(text)
    lowered = clean.casefold()
    if "тягач" in lowered and "полуприцеп" in lowered and "пневматическ" in lowered and "подвес" in lowered:
        return LOADING_TRACTOR_PNEUMO_LINE
    if "тягач" in lowered and "полуприцеп" in lowered:
        return LOADING_TRACTOR_LINE
    if "пневматическ" in lowered and "подвес" in lowered:
        return LOADING_PNEUMO_LINE
    if any(
        marker in lowered
        for marker in (
            "технически исправном и чистом состоянии",
            "технически исправном, чистом и сухом состоянии",
            "герметич",
            "целый, сухой тент",
            "стандартной обрешет",
            "без посторонних предмет",
        )
    ) and any(
        marker in lowered
        for marker in (
            "крепление грузовых мест",
            "механизмами, обеспечивающими крепление",
            "с механизмами для крепления",
        )
    ):
        return LOADING_TECH_FIXTURES_LINE
    if any(
        marker in lowered
        for marker in (
            "технически исправном и чистом состоянии",
            "технически исправном, чистом и сухом состоянии",
            "герметич",
            "целый, сухой тент",
            "стандартной обрешет",
            "без посторонних предмет",
        )
    ):
        return LOADING_TECH_LINE
    if any(
        marker in lowered
        for marker in (
            "крепление грузовых мест",
            "механизмами, обеспечивающими крепление",
            "с механизмами для крепления",
            "пол и стены транспортного средства должны иметь ровные поверхности",
            "пол и стены кузова",
        )
    ):
        return LOADING_FIXTURES_LINE
    if "грузовой отсек" in lowered and any(
        marker in lowered for marker in ("плотно закрываться", "попадания воды", "снега")
    ):
        return "Грузовой отсек должен плотно закрываться и защищать груз от попадания воды, снега и иных внешних воздействий"
    if any(
        marker in lowered
        for marker in (
            "видимых изменений конструкции",
            "усиливающих жесткость конструкции",
            "усилений бортов",
            "рефрижераторных установок",
        )
    ):
        return "Не допускаются видимые изменения конструкции ТС: усиления бортов/рамы, дополнительные металлические элементы, несогласованные рефрижераторные установки"
    if "подачу транспортных средств" in lowered and "технически исправном состоянии" in lowered:
        return (
            "Подача ТС в согласованные сроки, в технически исправном состоянии, "
            "с учетом требований к грузу и комплектом необходимых документов/разрешений"
        )
    return clean


def _merge_loading_requirements(items: list[str]) -> list[str]:
    result: list[str] = []
    has_combo = LOADING_TRACTOR_PNEUMO_LINE in items
    has_base = LOADING_TRACTOR_LINE in items
    has_pneumo = LOADING_PNEUMO_LINE in items
    has_specific_tech = LOADING_TECH_LINE in items
    has_specific_fixtures = LOADING_FIXTURES_LINE in items
    combo_inserted = False

    for item in items:
        if item == LOADING_TECH_FIXTURES_LINE and has_specific_tech and has_specific_fixtures:
            continue
        if item in {LOADING_TRACTOR_LINE, LOADING_PNEUMO_LINE, LOADING_TRACTOR_PNEUMO_LINE}:
            if has_combo:
                if item != LOADING_TRACTOR_PNEUMO_LINE or combo_inserted:
                    continue
                combo_inserted = True
                result.append(LOADING_TRACTOR_PNEUMO_LINE)
                continue
            if has_base and has_pneumo:
                if combo_inserted:
                    continue
                combo_inserted = True
                result.append(LOADING_TRACTOR_PNEUMO_LINE)
                continue
        result.append(item)

    return result


def _condense_special_condition(text: str) -> str:
    clean = _clean_block_text(text)
    lowered = clean.casefold()
    if "приложение gps системы контроля" in lowered and "установ" in lowered:
        return clean
    if "статусы" in lowered and "подтверд" in lowered and "выехал" in lowered and "выполнил" in lowered:
        return clean
    if "копию действующего страхового полиса" in lowered:
        return "Если страхование предусмотрено Поручением, Экспедитор предоставляет копию действующего страхового полиса"
    if "дополнительные виды услуг" in lowered:
        return "Дополнительные услуги оказываются только если они указаны в Поручении Экспедитору"
    if "охрана и страхование груза" in lowered:
        return "Охрана и страхование груза не входят в предмет договора, если иное не предусмотрено Поручением"
    return clean


def _condense_incident_action(text: str) -> str:
    clean = _clean_block_text(text)
    lowered = clean.casefold()
    if "информировать клиента" in lowered and "запрашивать инструкции" in lowered:
        return (
            "Информировать Клиента обо всех обстоятельствах, влияющих на стоимость и сроки доставки, "
            "и запрашивать инструкции по действиям с грузом. О задержках, авариях и других происшествиях "
            "сообщать в течение 1 рабочего дня"
        )
    if "принять все возможные меры" in lowered and "задержк" in lowered:
        return (
            "Принять все возможные меры для устранения задержки груза в пути. "
            "Предоставить документы, подтверждающие задержку и простой"
        )
    if "отклонени" in lowered and "маршрут" in lowered:
        return (
            "При отклонении от маршрута уведомить Клиента в течение 1 календарного дня "
            "и согласовать новый маршрут"
        )
    return clean


def _condense_penalty_item(text: str, group_key: str) -> str:
    clean = _clean_block_text(text)
    lowered = clean.casefold()

    if group_key == "delay":
        if (
            ("срыв подач" in lowered or "неподач" in lowered or "срок подачи транспортного средства" in lowered)
            and "50%" in lowered
            and "20%" in lowered
            and "стоимости услуг" in lowered
        ):
            return "За срыв подачи ТС: 50% для сетевых клиентов, 20% — для остальных"
        if "1000" in lowered and "простой" in lowered:
            if "за каждый час" in lowered or "за каждый полный час" in lowered:
                return "1000 рублей за каждый час сверхнормативного простоя при погрузке/разгрузке"
            if "сутк" in lowered or "за каждые полные сутки" in lowered:
                return "1000 рублей за каждые полные сутки сверхнормативного простоя при погрузке/разгрузке"
            return "1000 рублей за сверхнормативный простой при погрузке/разгрузке"
        if "500 рублей" in lowered and "задержк" in lowered and "подач" in lowered:
            return "500 рублей за каждый час задержки подачи транспортного средства"
        if "500 рублей" in lowered and "выгруз" in lowered and "20%" in lowered:
            return "500 рублей за каждый случай нарушения сроков вывоза ТС на выгрузку, но не более 20% от стоимости услуг"
        if "50%" in lowered and "стоимости услуг" in lowered and (
            "неподач" in lowered or "срыв подач" in lowered
        ):
            if "сетев" in lowered:
                return "50% от стоимости услуг за неподачу транспортного средства в срок для сетевых клиентов"
            return "50% от стоимости услуг за неподачу транспортного средства в срок"
        if "20%" in lowered and "стоимости услуг" in lowered and (
            "неподач" in lowered or "срыв подач" in lowered or "грузополучател" in lowered
        ):
            return "20% от стоимости услуг за неподачу транспортного средства в срок по прочим грузополучателям"
        if "10 %" in lowered or "10%" in lowered:
            return "10% от ставки при опоздании от 00 минут до одного часа"
        if "25%" in lowered:
            return "25% от ставки при опоздании от одного до двух часов"
        if "50%" in lowered:
            return "50% от ставки при опоздании от двух до четырёх часов"
        if "просрочки в доставке груза более чем на сутки" in lowered:
            return "20% от стоимости перевозки за каждые сутки опоздания, начиная со вторых суток"
        return ""

    if group_key == "trip_failure":
        if "срыв рейса" in lowered or "отказ от исполнения обязательств" in lowered:
            return "100% от ставки при отказе от исполнения подтвержденного Поручения. Срыв рейса — опоздание на загрузку более 4 часов"
        return clean

    if group_key == "documents":
        if "0,1%" in lowered or "0.1%" in lowered:
            return "0,1% от стоимости услуг за каждый день просрочки предоставления отчетных документов"
        if "500" in lowered and "документ" in lowered:
            return "500 рублей за каждый документ при несвоевременном предоставлении корректно оформленных документов"
        return clean

    if group_key == "gps":
        if "не предостав" in lowered and "100" in lowered:
            return "За не предоставление информации через Приложение GPS: 100 рублей за каждый случай"
        if "недостовер" in lowered and "100" in lowered:
            return "За предоставление недостоверной информации через Приложение GPS: 100 рублей за каждый случай"
        if "нарушает срок" in lowered and "50" in lowered:
            return "За нарушение срока предоставления информации через Приложение GPS: 50 рублей за каждый случай"
        return clean

    if group_key == "other":
        if "500 000" in lowered or "500000" in lowered:
            return "500 000 рублей ущерба при совершении действий, указанных в п. 2.1.13 договора"
        if "10%" in lowered and "суммы предложения" in lowered:
            return "10% от суммы предложения при отказе от подписания заявки"
        if "повреждения имущества клиента" in lowered or "складских комплексов" in lowered:
            return "Полное возмещение документально подтвержденного ущерба имуществу Клиента или нарушений правил поведения на складе в течение 10 календарных дней"
        if "штрафные санкции" in lowered and "10" in lowered and "претенз" in lowered:
            return "Штрафные санкции уплачиваются в течение 10 календарных дней с момента получения претензии и подтверждающих документов"
        if any(marker in lowered for marker in ("нагрузки на ось", "допустимой массы", "в полном объеме")):
            return "Экспедитор компенсирует в полном объеме штрафы за превышение допустимой массы ТС и/или нагрузки на ось"
        return clean

    return clean


def _condense_document_requirement(text: str) -> str:
    clean = _clean_block_text(text)
    lowered = clean.casefold()
    if "доверен" in lowered and any(
        marker in lowered for marker in ("получени", "груз", "водител", "факсим", "электрон", "оригинал")
    ):
        return DOCUMENT_TRUST_REQUIREMENT_LINE
    return clean


def _render_status_frequency(data: CustomerInstructionData) -> str:
    frequency = _clean_block_text(data.status_informing.frequency)
    lowered = frequency.casefold()
    fact_pool = [
        *data.special_conditions,
        *data.driver_briefing,
        *data.incident_actions,
        frequency,
    ]
    facts_text = " ".join(_clean_block_text(item) for item in fact_pool if _clean_block_text(item)).casefold()
    has_gps_statuses = (
        "gps" in facts_text
        and any(marker in facts_text for marker in ("подтвердил", "выехал", "выполнил", "30 минут"))
    )
    has_delay_notice = (
        "1 рабочего дня" in facts_text
        or "1 (одного) рабочего дня" in facts_text
        or "отклонени" in facts_text
        or "авари" in facts_text
        or "задерж" in facts_text
    )

    if has_gps_statuses:
        parts = [
            "По каждой заявке: статусы Подтвердил / Выехал / Выполнил проставляются в приложении GPS не позднее 30 минут после фактического действия"
        ]
        if has_delay_notice:
            parts.append("О задержках, авариях и отклонениях от маршрута сообщать в течение 1 рабочего дня")
        return ". ".join(parts)

    if not frequency or lowered in {"незамедлительно", "сразу", "по факту изменений"}:
        if has_delay_notice:
            return "По факту изменений; о задержках, авариях и отклонениях от маршрута сообщать в течение 1 рабочего дня"
    return frequency


def _condense_edo_workflow(text: str) -> str:
    clean = _clean_block_text(text)
    lowered = clean.casefold()
    if (
        "электронного документооборота" in lowered
        and "10" in lowered
        and "штраф" in lowered
        and "документ" in lowered
    ):
        return (
            "При требовании Клиента о корректировке/исправлении документов Экспедитор "
            "вносит исправления в течение 10 дней. Иначе Клиент вправе отказаться от договора "
            "и/или взыскать 500 рублей за каждый документ"
        )
    return clean


def _build_incident_blocks(items: list[str]) -> list[ParagraphBlock]:
    if not items:
        return [ParagraphBlock(MISSING_MARKER)]
    if len(items) == 1:
        return [ParagraphBlock(_ensure_sentence(items[0]), style_name=BODY_STYLE)]
    return [
        ParagraphBlock(_ensure_sentence(item), style_name=LIST_STYLE, bullet=True)
        for item in items
    ]


def _build_payment_package_lines(data: CustomerInstructionData) -> list[str]:
    general_items = [
        item
        for item in data.payment_document_package
        if not item.lower().startswith("автомобильные перевозки:")
        and not item.lower().startswith("железнодорожный транспорт:")
    ]

    lines: list[str] = []
    if general_items:
        lines.extend(_build_bulleted_lines(general_items))
    if data.payment_document_package_auto:
        if lines:
            lines.append("")
        lines.append("Автомобильные перевозки:")
        lines.extend(_build_bulleted_lines(data.payment_document_package_auto))
    if data.payment_document_package_rail:
        if lines:
            lines.append("")
        lines.append("Железнодорожный транспорт:")
        lines.extend(_build_bulleted_lines(data.payment_document_package_rail))

    if not lines and data.payment_document_package:
        lines.extend(_build_bulleted_lines(data.payment_document_package))

    return lines or [MISSING_MARKER]


def _build_bulleted_lines(items: list[str]) -> list[str]:
    normalized_items = [_clean_block_text(item) for item in items if _clean_block_text(item)]
    if not normalized_items:
        return [MISSING_MARKER]
    return [f"• {_ensure_sentence(item)}" for item in normalized_items]


def _build_payment_package_blocks(data: CustomerInstructionData) -> list[ParagraphBlock]:
    blocks: list[ParagraphBlock] = []
    general_items = [
        item
        for item in data.payment_document_package
        if not item.lower().startswith("автомобильные перевозки:")
        and not item.lower().startswith("железнодорожный транспорт:")
    ]

    for item in general_items:
        blocks.append(_table_bullet_block(item))

    if data.payment_document_package_auto:
        if blocks:
            blocks.append(_table_text_block(""))
        blocks.append(_table_text_block("Автомобильные перевозки:", bold=True))
        for item in data.payment_document_package_auto:
            blocks.append(_table_bullet_block(item))

    if data.payment_document_package_rail:
        if blocks:
            blocks.append(_table_text_block(""))
        blocks.append(_table_text_block("Железнодорожный транспорт:", bold=True))
        for item in data.payment_document_package_rail:
            blocks.append(_table_bullet_block(item))

    return blocks or [_table_text_block(MISSING_MARKER)]


def _build_document_requirement_blocks(items: list[str]) -> list[ParagraphBlock]:
    cleaned_items = _prepare_document_requirements(items)
    if not cleaned_items:
        return [_table_text_block(MISSING_MARKER)]
    return [_table_bullet_block(item) for item in cleaned_items]


def _build_payment_blocks(data: CustomerInstructionData) -> list[ParagraphBlock]:
    blocks: list[ParagraphBlock] = []
    if data.payment_term:
        blocks.append(ParagraphBlock(_ensure_sentence(data.payment_term), style_name=BODY_STYLE))
    if data.payment_hold_condition:
        blocks.append(ParagraphBlock(_ensure_sentence(data.payment_hold_condition), style_name=BODY_STYLE))
    if data.tax_change_notification:
        blocks.append(ParagraphBlock(_ensure_sentence(data.tax_change_notification), style_name=BODY_STYLE))
    if data.client_payment_delay_penalty:
        blocks.append(
            ParagraphBlock(
                _ensure_sentence(data.client_payment_delay_penalty),
                style_name=BODY_STYLE,
                penalty_highlight=True,
            )
        )
    if data.copies_followed_by_originals:
        blocks.append(ParagraphBlock(_ensure_sentence(data.copies_followed_by_originals), style_name=BODY_STYLE))
    if not blocks:
        blocks.append(ParagraphBlock(MISSING_MARKER, style_name=BODY_STYLE))
    return blocks


def _normalize_platform_name(row: PlatformRow) -> str:
    platform_name = _clean_block_text(row.platform_name)
    if platform_name:
        return platform_name

    bidding_rules = _clean_block_text(row.bidding_rules).lower()
    if "почт" in bidding_rules or "e-mail" in bidding_rules or "email" in bidding_rules:
        return "Рассылка на почту"
    return MISSING_MARKER


def _select_communication_row(
    table,
    candidate: ContactRow,
    template_slots: list[int],
    assigned_rows: set[int],
) -> int | None:
    best_index: int | None = None
    best_score = -1

    for row_index in template_slots:
        if row_index in assigned_rows:
            continue

        slot_cells = table.rows[row_index].cells
        slot_text = " ".join(cell.text for cell in slot_cells)
        score = _communication_match_score(slot_text, candidate)
        if score > best_score:
            best_score = score
            best_index = row_index

    if best_score > 0:
        return best_index

    for row_index in template_slots:
        if row_index not in assigned_rows:
            return row_index
    return None


def _communication_match_score(slot_text: str, candidate: ContactRow) -> int:
    slot = slot_text.casefold()
    haystack = " ".join(
        value
        for value in (
            candidate.full_name,
            candidate.role,
            candidate.contacts,
            candidate.responsibility,
        )
        if value
    ).casefold()

    rules = (
        ("заяв", 4),
        ("логист", 4),
        ("сервис", 3),
        ("руковод", 2),
        ("договор", 4),
        ("закуп", 4),
        ("бух", 5),
        ("документооборот", 5),
        ("претенз", 5),
        ("финанс", 4),
    )

    score = 0
    for token, weight in rules:
        if token in slot and token in haystack:
            score += weight

    if "по всем вопросам" in haystack and ("логист" in slot or "руковод" in slot):
        score += 2
    if "@" in candidate.contacts and "общий адрес" in slot:
        score += 1
    if "тендер" in haystack and "руковод" in slot:
        score += 1
    return score


def _expand_penalty_item(item: str) -> list[str]:
    cleaned = _clean_block_text(item)
    if not cleaned:
        return []
    if ";" in cleaned and sum("%" in part for part in cleaned.split(";")) >= 2:
        return [part.strip() for part in cleaned.split(";") if part.strip()]
    return [cleaned]


def _looks_like_compact_requirement(text: str) -> bool:
    lowered = text.casefold()
    if not lowered:
        return False
    if text.endswith(";"):
        return True
    if text[:1].islower():
        return True

    sentence_markers = (
        "экспедитор",
        "клиент",
        "по прибытии",
        "в случае",
        "водител",
        "информир",
        "предостав",
        "обязательно",
        "уведомлен",
        "оплата",
        "штраф",
    )
    return len(text) < 170 and not any(marker in lowered for marker in sentence_markers)


def _is_intro_like(existing_text: str, intro_text: str) -> bool:
    return existing_text.casefold().startswith(intro_text.casefold()[:18])


def _ensure_sentence(text: str) -> str:
    cleaned = _clean_block_text(text)
    if not cleaned:
        return MISSING_MARKER
    if cleaned.endswith((".", ";", "!", "?", ":")):
        return cleaned
    return f"{cleaned}."


def _ensure_compact_requirement(text: str, is_last: bool) -> str:
    cleaned = _clean_block_text(text).rstrip(".;")
    if not cleaned:
        return MISSING_MARKER
    suffix = "." if is_last else ";"
    return f"{cleaned}{suffix}"


def _clean_block_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _format_markdown_blocks(blocks: list[ParagraphBlock]) -> str:
    if not blocks:
        return f"- {MISSING_MARKER}"

    lines: list[str] = []
    for block in blocks:
        if block.bullet or block.style_name == LIST_STYLE:
            lines.append(f"- {block.text}")
        else:
            lines.append(block.text)
    return "\n".join(lines)


def _format_indented_lines(lines: list[str]) -> str:
    normalized = lines or [MISSING_MARKER]
    return "\n".join(f"  {line}" if line else "" for line in normalized)


def _format_optional_list(items: list[str], empty_message: str) -> str:
    if not items:
        return f"- {empty_message}"
    return "\n".join(f"- {item}" for item in items)


def _format_contacts(data: CustomerInstructionData) -> str:
    if not data.communication_rows:
        return f"- {MISSING_MARKER}"
    return "\n".join(
        f"- {_contact_display_name(row)} | {_contact_display_role(row)} | {row.contacts or MISSING_MARKER} | {row.responsibility or MISSING_MARKER}"
        for row in data.communication_rows
    )


def _contact_display_name(row: ContactRow) -> str:
    if row.full_name:
        return row.full_name
    if row.role and row.role != "Общий адрес клиента":
        return row.role
    if row.contacts:
        return "Общий адрес клиента"
    return MISSING_MARKER


def _contact_display_role(row: ContactRow) -> str:
    if row.role == "Общий адрес клиента":
        return ""
    if row.role:
        return row.role
    if row.contacts:
        return ""
    return ""


def _append_payment_section(document: Document, data: CustomerInstructionData) -> None:
    if _find_paragraph(document, "12. Условия оплаты") is not None:
        return

    document.add_paragraph("")
    heading = document.add_paragraph()
    heading.style = "Normal"
    heading_run = heading.add_run("12. Условия оплаты")
    heading_run.bold = True
    heading_run.italic = True
    heading_run.font.name = "Times New Roman"
    heading_run.font.size = Pt(14)

    for block in _build_payment_blocks(data):
        paragraph = document.add_paragraph()
        _populate_paragraph(paragraph, block)


def _replace_heading_text(
    document: Document,
    current_text: str,
    new_text: str,
) -> None:
    paragraph = _find_paragraph(document, current_text)
    if paragraph is None:
        return

    _clear_paragraph_content(paragraph)
    paragraph.style = "Normal"
    run = paragraph.add_run(new_text)
    run.bold = True
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def _apply_section_spacing(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if SECTION_HEADING_RE.match(text):
            is_inner_numbered_block = getattr(paragraph.style, "name", "") == BODY_STYLE
            paragraph.paragraph_format.space_before = Pt(8 if is_inner_numbered_block else 18)
            paragraph.paragraph_format.space_after = Pt(4 if is_inner_numbered_block else 8)
            paragraph.paragraph_format.keep_with_next = True
        elif text in {
            "Гарантированные заявки",
            "Спотовые заявки",
            "Информирование Клиента о статусе перевозки",
            "Информирование клиента о статусе перевозки",
        }:
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(6)


def _apply_accent_styling(document: Document) -> None:
    accent_labels = {
        "Гарантированные заявки",
        "Спотовые заявки",
    }
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in accent_labels:
            for run in paragraph.runs:
                run.font.color.rgb = ACCENT_BLUE


def _normalize_cell_block(item, style_name: str) -> ParagraphBlock:
    if isinstance(item, ParagraphBlock):
        block = item
        if not block.style_name:
            block.style_name = style_name
        return block
    return ParagraphBlock(str(item), style_name=style_name)


def _table_text_block(text: str, bold: bool = False) -> ParagraphBlock:
    return ParagraphBlock(text, style_name=TABLE_STYLE, bold=bold)


def _table_bullet_block(text: str) -> ParagraphBlock:
    cleaned = _clean_block_text(text)
    if cleaned.startswith("• "):
        cleaned = cleaned[2:].strip()
    return ParagraphBlock(
        _ensure_sentence(cleaned),
        style_name=TABLE_STYLE,
        bullet=True,
    )


def _segment_penalty_text(text: str) -> list[tuple[str, bool]]:
    spans = _find_penalty_highlight_spans(text)
    if not spans:
        return [(text, False)]

    parts: list[tuple[str, bool]] = []
    cursor = 0
    for start, end in spans:
        if start > cursor:
            parts.append((text[cursor:start], False))
        parts.append((text[start:end], True))
        cursor = end
    if cursor < len(text):
        parts.append((text[cursor:], False))
    return parts


def _find_penalty_highlight_spans(text: str) -> list[tuple[int, int]]:
    spans: list[tuple[int, int]] = []

    for pattern in PENALTY_HIGHLIGHT_PATTERNS:
        for match in pattern.finditer(text):
            spans.append(match.span())

    for match in re.finditer(r"\b\d+\s?%\b", text):
        spans.append(match.span())
    for match in re.finditer(
        r"\b\d[\d\s]*(?:,\d+)?\s*руб(?:л(?:ей|я)|\.)\b",
        text,
        flags=re.IGNORECASE,
    ):
        spans.append(match.span())
    for match in re.finditer(r"\b\d+\s+час(?:а|ов)?\b", text, flags=re.IGNORECASE):
        spans.append(match.span())

    if not spans:
        return []

    spans.sort()
    merged: list[tuple[int, int]] = [spans[0]]
    for start, end in spans[1:]:
        last_start, last_end = merged[-1]
        if start <= last_end:
            merged[-1] = (last_start, max(last_end, end))
        else:
            merged.append((start, end))
    return merged
