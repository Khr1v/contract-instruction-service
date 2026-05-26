from __future__ import annotations

from docx import Document

from app.documents.docx_extractor import DOCXExtractor
from app.llm.schemas import SourceFormat


def test_docx_extractor_reads_paragraphs_and_tables(tmp_path):
    path = tmp_path / "contract.docx"
    document = Document()
    document.add_paragraph("Договор оказания услуг № 1")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Условие"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "Оплата"
    table.cell(1, 1).text = "10 дней"
    document.save(path)

    canonical = DOCXExtractor().extract(path, source_file_id="doc-1", filename="contract.docx")

    assert canonical.source_format == SourceFormat.DOCX
    assert "Договор оказания услуг" in canonical.document_text
    assert "| Условие | Значение |" in canonical.document_text
    assert canonical.tables
    assert canonical.quality_score == 1.0

