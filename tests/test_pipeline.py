from __future__ import annotations

import orjson
import pytest
from docx import Document

from app.config import Settings
from app.services.contract_pipeline import ContractPipeline


class FakeLLM:
    def generate_text(
        self,
        *,
        instructions: str,
        input: str,
        temperature: float | None = None,
        max_output_tokens: int | None = None,
        model_name: str | None = None,
    ) -> str:
        if "данные для заполнения Word-шаблона" in instructions:
            return orjson.dumps(
                {
                    "client_name": "ООО Клиент",
                    "contract_legal_entity": "ООО Клиент",
                    "generated_date": "01.01.2026",
                    "work_format": "Спотовые заявки",
                    "guaranteed_application_rules": ["Уточнить у человека"],
                    "spot_application_rules": ["Заявки по договоренности"],
                    "communication_rows": [
                        {
                            "full_name": "Уточнить у человека",
                            "role": "Логистика",
                            "contacts": "Уточнить у человека",
                            "responsibility": "Уточнить у человека",
                        }
                    ],
                    "loading_requirements": ["Уточнить у человека"],
                    "unloading_requirements": ["Уточнить у человека"],
                    "special_conditions": ["Уточнить у человека"],
                    "driver_briefing": ["Проверить документы"],
                    "penalties": ["Пени 0,1%"],
                    "incident_actions": ["Сообщить логисту"],
                    "status_informing": {"is_required": "Да", "frequency": "По заявкам", "channels": []},
                    "payment_document_package": ["Акт", "Счет"],
                    "document_format_requirements": ["Корректные документы"],
                    "edo_workflow": "ЭДО Диадок",
                    "payment_term": "Оплата в течение 10 дней",
                    "open_questions": ["Уточнить у человека: нет текста всех приложений"],
                    "extraction_notes": [],
                }
            ).decode("utf-8")
        if "Верни только валидный JSON" in instructions:
            return orjson.dumps(
                {
                    "contract_type": self.fact("Договор оказания услуг"),
                    "contract_number": self.fact("1"),
                    "contract_date": self.fact("01.01.2026"),
                    "parties": {
                        "client": self.fact("ООО Клиент"),
                        "contractor": self.fact("ООО Исполнитель"),
                    },
                    "subject": self.fact("Оказание услуг"),
                    "service_description": self.fact("Логистические услуги"),
                    "term_and_duration": self.fact("До 31.12.2026"),
                    "payment_terms": self.fact("Оплата в течение 10 дней"),
                    "document_flow": self.fact("Акт и счет"),
                    "edo_terms": self.fact("ЭДО Диадок"),
                    "delivery_or_service_terms": self.fact("По заявкам"),
                    "responsibilities": [self.fact("Стороны несут ответственность")],
                    "penalties": [self.fact("Пени 0,1%")],
                    "fines": [],
                    "deadlines": [self.fact("10 дней на оплату")],
                    "reporting_requirements": [],
                    "acceptance_procedure": self.fact("Подписание акта"),
                    "termination_terms": self.fact("По соглашению сторон"),
                    "confidentiality": self.fact("Коммерческая тайна"),
                    "special_conditions": [],
                    "appendices": [self.fact("Приложение №1")],
                    "requisites": self.fact("Реквизиты сторон"),
                    "missing_information": ["Уточнить у человека: нет текста всех приложений"],
                    "risk_flags": [],
                    "human_review_required": True,
                }
            ).decode("utf-8")
        if "Ты проверяешь готовую инструкцию" in instructions:
            return '{"is_valid": true, "problems": [], "missing_sections": [], "critical_warnings": [], "recommendation": "approve"}'
        return """# ИНСТРУКЦИЯ ПО РАБОТЕ
С КЛИЕНТОМ
Клиент: ООО Клиент
Юридическое лицо по договору: ООО Исполнитель
# 1. Форма работы
Уточнить у человека: нет текста всех приложений.
Гарантированные заявки
| Направление перевозки | Стоимость<br>перевозки | Кол-во ТС | Особые условия |
| --- | --- | --- | --- |
| Уточнить у человека | Уточнить у человека | Уточнить у человека | Уточнить у человека |
Спотовые заявки
| Название площадки | Логин и<br>пароль | Порядок проведения торгов и выбора победителей | Ссылка на инструкцию по<br>работе с площадкой |
| --- | --- | --- | --- |
| Уточнить у человека | Уточнить у человека | Уточнить у человека | Уточнить у человека |
# 2. Матрица коммуникаций Клиента
| ФИО сотрудника | Должность | Контакты | Зона ответственности |
| --- | --- | --- | --- |
| Уточнить у человека | Уточнить у человека | Уточнить у человека | Уточнить у человека |
# 5. Требования на погрузке
Уточнить у человека.
# 6. Требования на выгрузке
Уточнить у человека.
# 7. Особые условия
Уточнить у человека.
# 8. Инструктаж для водителя
Уточнить у человека.
# 9. Штрафы
Пени 0,1%.
# 10. Действия логиста при проблемных ситуациях
Уточнить у человека.
# 11. Информирование Клиента о статусе перевозки
| Требуется регулярное информирование | Да / Нет |
| --- | --- |
| Частота информирования о статусе перевозки | Уточнить у человека |
# 11. Документооборот
| Комплект документов для оплаты: | Акт и счет |
| --- | --- |
| Работа через ЭДО | ЭДО Диадок |
Уточнить у человека: нет текста всех приложений."""

    def fact(self, value: str) -> dict[str, object]:
        return {
            "value": value,
            "source_ref": "section_1",
            "source_quote": value,
            "confidence": "high",
            "needs_human_review": False,
        }


class FakeRAG:
    def get_instruction_template(self) -> str:
        return "# ИНСТРУКЦИЯ ПО РАБОТЕ\nС КЛИЕНТОМ\n# 1. Форма работы\n"

    def get_relevant_rules(self, query: str) -> list[str]:
        return ["Если данных нет, писать: Уточнить у человека."]

    def get_relevant_examples(self, query: str) -> list[str]:
        return []


@pytest.mark.asyncio
async def test_pipeline_happy_path_docx(tmp_path):
    data_dir = tmp_path / "data"
    settings = Settings(
        _env_file=None,
        DATA_DIR=data_dir,
        UPLOADS_DIR=data_dir / "uploads",
        PROCESSED_DIR=data_dir / "processed",
        VECTORSTORE_DIR=data_dir / "vectorstore",
        TEMPLATES_DIR=data_dir / "templates",
        SQLITE_DB_PATH=data_dir / "app.db",
        YANDEX_CLOUD_API_KEY="fake",
    )
    settings.ensure_directories()

    document_path = tmp_path / "contract.docx"
    doc = Document()
    doc.add_paragraph("Договор оказания услуг № 1 от 01.01.2026")
    doc.add_paragraph("Оплата в течение 10 дней. Пени 0,1%.")
    doc.save(document_path)

    pipeline = ContractPipeline(settings=settings, llm_client=FakeLLM(), rag_agent=FakeRAG())
    result = await pipeline.process_contract(
        file_path=str(document_path),
        original_filename="contract.docx",
        external_user_id="user-1",
        source_channel="api",
    )

    assert result.status == "completed"
    assert result.instruction_path is not None
    assert result.facts_json_path is not None
    assert result.human_review_required
    assert "9. Штрафы" in (result.instruction_markdown or "")
